import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_hyperlink(paragraph, url, text, color="1877F2"):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create hyperlink element with proper namespaces
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        f'r:id="{r_id}"/>'
    )
    
    new_run = parse_xml(
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:rPr><w:color w:val="{color}"/><w:u w:val="single"/></w:rPr>'
        f'<w:t>{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_docx(groups, output_path="Danh_sach_151_nhom_Facebook.docx"):
    doc = docx.Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("DANH SÁCH CÁC NHÓM FACEBOOK ĐÃ THAM GIA")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(24, 119, 242) # Facebook Blue

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Tổng cộng: {len(groups)} nhóm Facebook")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph() # spacing

    # Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "STT"
    hdr_cells[1].text = "Tên Nhóm Facebook"
    hdr_cells[2].text = "Link Nhóm"

    col_widths = [Inches(0.6), Inches(3.2), Inches(3.2)]
    
    # Style Header
    for i, cell in enumerate(hdr_cells):
        cell.width = col_widths[i]
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1877F2"/>'))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = "Arial"
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for idx, g in enumerate(groups, 1):
        row_cells = table.add_row().cells
        
        # Row zebra background for alternating rows
        bg_color = "F0F2F5" if idx % 2 == 0 else "FFFFFF"
        for i, cell in enumerate(row_cells):
            cell.width = col_widths[i]
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>'))

        # STT
        p0 = row_cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(idx))
        r0.font.name = "Arial"
        r0.font.size = Pt(9.5)

        # Name
        p1 = row_cells[1].paragraphs[0]
        name = str(g.get("name", "Nhóm Facebook")).strip()
        r1 = p1.add_run(name)
        r1.font.name = "Arial"
        r1.font.size = Pt(9.5)
        r1.font.bold = True

        # Link
        p2 = row_cells[2].paragraphs[0]
        url = str(g.get("url", "")).strip()
        if url:
            add_hyperlink(p2, url, url, color="1877F2")
        else:
            p2.add_run("N/A")

    doc.save(output_path)
    print(f"[OK] Docx saved: {output_path} with {len(groups)} groups.")

if __name__ == "__main__":
    import sys, json
    if len(sys.argv) > 1:
        json_file = sys.argv[1]
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        out = sys.argv[2] if len(sys.argv) > 2 else "d:/HUDI/n8n/Danh_sach_151_nhom_Facebook.docx"
        create_docx(data, out)
    else:
        sample = [
            {"name": "Cộng Đồng Thiết Kế Website, Landing Page, App", "url": "https://www.facebook.com/groups/123456789/"},
            {"name": "Web Developer In USA", "url": "https://www.facebook.com/groups/webdevusa/"}
        ]
        create_docx(sample, "d:/HUDI/n8n/test_groups.docx")
